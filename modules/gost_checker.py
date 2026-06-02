from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import zipfile
import shutil
import os

def repair_docx(input_path, output_path):
    with zipfile.ZipFile(input_path, "r") as zin:
        with zipfile.ZipFile(output_path, "w") as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename.endswith(".rels"):
                    text = data.decode("utf-8", errors="ignore")

                    text = re.sub(
                        r'<Relationship[^>]+Target="#[^"]+"[^>]*/>',
                        '',
                        text
                    )

                    data = text.encode("utf-8")

                zout.writestr(item, data)

def is_bold_paragraph(paragraph):

    for run in paragraph.runs:
        if run.bold is True:
            return True

    try:
        if paragraph.style and paragraph.style.font.bold is True:
            return True
    except:
        pass

    return False

def get_paragraph_alignment(paragraph):

    if paragraph.alignment is not None:
        return paragraph.alignment

    try:
        if paragraph.style and paragraph.style.paragraph_format.alignment is not None:
            return paragraph.style.paragraph_format.alignment
    except:
        pass

    return None


def is_centered(paragraph):

    alignment = get_paragraph_alignment(paragraph)

    if alignment is None:
        # Если Word не отдаёт выравнивание напрямую,
        # не считаем это ошибкой.
        return True

    return alignment == WD_PARAGRAPH_ALIGNMENT.CENTER


def is_justified(paragraph):

    alignment = get_paragraph_alignment(paragraph)

    if alignment is None:
        # Часто основной стиль уже задан по ширине,
        # но python-docx возвращает None.
        # Чтобы не ловить ложные ошибки, пропускаем.
        return True

    return alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY

def extract_gost_rules(method_text=""):

    method_lower = method_text.lower() if method_text else ""

    rules = {
        "main_heading_center": True,
        "main_heading_bold": True,
        "main_heading_caps": True,

        # По умолчанию не ругаем подразделы за центр,
        # потому что методички часто разрешают центрирование.
        "subheading_allow_center": True,

        "subheading_bold": True,

        "font_name": "Times New Roman",
        "font_size": 14,
        "line_spacing_min": 1.4,
        "line_spacing_max": 1.6,
        "indent_min": 10,
        "indent_max": 20
    }

    # Если методичка явно говорит, что заголовки/подразделы по центру
    if (
        "заголовки располагаются по центру" in method_lower
        or "заголовки выравниваются по центру" in method_lower
        or "выравнивание заголовков по центру" in method_lower
        or "заголовки следует располагать по центру" in method_lower
        or "заголовки оформляются по центру" in method_lower
    ):
        rules["subheading_allow_center"] = True

    # Если методичка явно говорит, что подразделы НЕ по центру
    if (
        "подразделы выравниваются по левому краю" in method_lower
        or "заголовки подразделов выравниваются по левому краю" in method_lower
        or "подразделы располагаются с абзацного отступа" in method_lower
    ):
        rules["subheading_allow_center"] = False

    return rules

def check_tables_and_figures(doc):

    tables_ok = True
    figures_ok = True

    table_errors = []
    figure_errors = []

    table_numbers = []
    figure_numbers = []

    paragraphs_text = [
        paragraph.text.strip()
        for paragraph in doc.paragraphs
        if paragraph.text.strip()
    ]

    full_text = "\n".join(paragraphs_text)

    # =====================
    # ПОИСК ПОДПИСЕЙ ТАБЛИЦ
    # =====================

    table_matches = re.findall(
        r"(таблица\s+(\d+(?:\.\d+)?)[^\n]*)",
        full_text,
        flags=re.IGNORECASE
    )

    for full_caption, number in table_matches:
        table_numbers.append(number)

        if not re.search(
            r"таблица\s+\d+(?:\.\d+)?\s*[–—-]\s*\S+",
            full_caption,
            flags=re.IGNORECASE
        ):
            tables_ok = False
            table_errors.append(
                f"Некорректная подпись таблицы: {full_caption}"
            )

    # Если таблицы в документе есть, а подписей нет
    if len(doc.tables) > 0 and not table_matches:
        tables_ok = False
        table_errors.append(
            "В документе есть таблицы, но не найдены подписи вида 'Таблица 1 — Название'."
        )

    # Проверка последовательности номеров таблиц
    if table_numbers:
        # Проверяем только дубли, а не строгую последовательность,
        # потому что допускается нумерация по главам: 2.1, 2.2 и т.д.
        if len(table_numbers) != len(set(table_numbers)):
            tables_ok = False
            table_errors.append(
                "Обнаружены повторяющиеся номера таблиц."
            )

    # =====================
    # ПОИСК ПОДПИСЕЙ РИСУНКОВ
    # =====================

    figure_matches = re.findall(
        r"(рисунок\s+(\d+(?:\.\d+)?)[^\n]*)",
        full_text,
        flags=re.IGNORECASE
    )

    for full_caption, number in figure_matches:
        figure_numbers.append(number)

        if not re.search(
            r"рисунок\s+\d+(?:\.\d+)?\s*[–—-]\s*\S+",
            full_caption,
            flags=re.IGNORECASE
        ):
            figures_ok = False
            figure_errors.append(
                f"Некорректная подпись рисунка: {full_caption}"
            )

    # Проверка последовательности номеров рисунков
    if figure_numbers:
        if len(figure_numbers) != len(set(figure_numbers)):
            figures_ok = False
            figure_errors.append(
                "Обнаружены повторяющиеся номера рисунков."
            )

    return {
        "tables": {
            "status": tables_ok,
            "errors": table_errors,
            "found": len(table_matches),
            "real_tables": len(doc.tables)
        },
        "figures": {
            "status": figures_ok,
            "errors": figure_errors,
            "found": len(figure_matches)
        }
    }

def check_gost(file_path, method_text=""):

    rules = extract_gost_rules(method_text)

    try:
        doc = Document(file_path)

    except Exception:
        repaired_path = "temp_repaired.docx"

        repair_docx(file_path, repaired_path)

        doc = Document(repaired_path)

    start_checking = False
    in_literature = False
    in_contents = False

    font_ok = True
    size_ok = True
    spacing_ok = True
    alignment_ok = True
    indent_ok = True
    headings_ok = True

    heading_errors = []

    font_errors = set()
    size_errors = set()

    # =====================
    # ПАТТЕРНЫ ЗАГОЛОВКОВ
    # =====================
    major_heading_patterns = [
        r"^введение$",
        r"^заключение$",
        r"^содержание$",
        r"^список использованных источников$",
        r"^список литературы$",
        r"^приложения?$"
    ]

    # Только настоящие главы (более строгий паттерн)
    chapter_patterns = [
        r"^глава\s+\d",           # Глава 1
        r"^\d+\.\d+\s",           # 1.2 Название главы
        r"^\d+\.\s+[А-Я]"         # 1. Название с большой буквы (чтобы не ловить список литературы)
    ]

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        text_lower = text.lower()

        # =====================
        # ПРОПУСК СОДЕРЖАНИЯ
        # =====================

        if text_lower in ["содержание", "оглавление"]:
            in_contents = True
            start_checking = False
            continue

        if in_contents:
            # В содержании обычно идут ВВЕДЕНИЕ, главы, ЗАКЛЮЧЕНИЕ, список источников.
            # Настоящий текст начинается после содержания, когда появляется первое содержательное название главы.
            if (
                re.match(r"^глава\s+1", text_lower)
                or re.match(r"^1\.\s+[а-яa-z]", text_lower)
                or re.match(r"^1\.1\s+[а-яa-z]", text_lower)
            ):
                in_contents = False
                start_checking = True
            else:
                continue

        # =====================
        # ОПРЕДЕЛЯЕМ РАЗДЕЛЫ
        # =====================
        if "список использованных источников" in text_lower or "список литературы" in text_lower:
            in_literature = True

        # =====================
        # ПРОВЕРКА ЗАГОЛОВКОВ
        # =====================
        is_heading = False

        # Главные заголовки (ВВЕДЕНИЕ, СПИСОК ЛИТЕРАТУРЫ и т.д.)
        for pattern in major_heading_patterns:
            if re.match(pattern, text_lower):
                is_heading = True
                # CAPS + Центр + Жирный
                if text != text.upper():
                    headings_ok = False
                    heading_errors.append(f"Главный заголовок должен быть CAPSLOCK: {text}")

                if not is_centered(paragraph):
                    headings_ok = False
                    heading_errors.append(f"Главный заголовок должен быть по центру: {text}")

                bold_found = is_bold_paragraph(paragraph)
                if not bold_found:
                    headings_ok = False
                    heading_errors.append(f"Главный заголовок должен быть жирным: {text}")
                break

        # Главы и подразделы (только если не в списке литературы)
        if not is_heading and not in_literature:
            for pattern in chapter_patterns:
                if re.match(pattern, text_lower):
                    is_heading = True
                    # Только жирный + не по центру
                    bold_found = is_bold_paragraph(paragraph)
                    if not bold_found:
                        headings_ok = False
                        heading_errors.append(f"Глава/подраздел должен быть жирным: {text}")

                    # Главы и подразделы (только если не в списке литературы)
                    if not is_heading and not in_literature:
                        for pattern in chapter_patterns:
                            if re.match(pattern, text_lower):
                                is_heading = True

                                bold_found = is_bold_paragraph(paragraph)

                                if not bold_found:
                                    headings_ok = False
                                    heading_errors.append(
                                        f"Глава/подраздел должен быть жирным: {text}"
                                    )

                                break

        # =====================
        # ПРОВЕРКА ОСНОВНОГО ТЕКСТА
        # =====================
        if not start_checking or in_literature or is_heading:
            continue   # Пропускаем титульный лист, список литературы и сами заголовки

        # Шрифт и размер
        for run in paragraph.runs:
            font = run.font
            if font.name and font.name != "Times New Roman":
                font_ok = False
                font_errors.add(font.name)

            if font.size:
                size = round(font.size.pt)
                if size != 14:
                    size_ok = False
                    size_errors.add(size)

        # Абзац
        fmt = paragraph.paragraph_format

        if fmt.line_spacing:
            try:
                spacing = float(fmt.line_spacing)
                if spacing < 1.4 or spacing > 1.6:
                    spacing_ok = False
            except:
                pass

        if fmt.first_line_indent:
            try:
                indent = fmt.first_line_indent.mm
                if indent < 10 or indent > 20:
                    indent_ok = False
            except:
                pass

        if not is_justified(paragraph):
            alignment_ok = False

    tables_figures_result = check_tables_and_figures(doc)

    return {
        "font": {"status": font_ok, "errors": list(font_errors)},
        "size": {"status": size_ok, "errors": list(size_errors)},
        "spacing": {"status": spacing_ok},
        "alignment": {"status": alignment_ok},
        "indent": {"status": indent_ok},
        "headings": {"status": headings_ok, "errors": heading_errors},
        "tables": tables_figures_result["tables"],
        "figures": tables_figures_result["figures"]
    }